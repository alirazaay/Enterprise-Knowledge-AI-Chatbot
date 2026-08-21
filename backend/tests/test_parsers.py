from pathlib import Path

import pymupdf as fitz
import pytest
from docx import Document as DocxDocument

from app.services.docx_parser import DocxParser
from app.services.document_parser import ParserError
from app.services.pdf_parser import PdfParser


def make_pdf(path: Path) -> None:
    document = fitz.open()
    document.new_page().insert_text((72, 72), "Heading\nFirst page policy text.")
    document.new_page().insert_text((72, 72), "Second page content.")
    document.save(path)
    document.close()


def make_docx(path: Path) -> None:
    document = DocxDocument()
    document.add_heading("Leave Policy", level=1)
    document.add_paragraph("Employees may request annual leave through their manager.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Department"
    table.rows[0].cells[1].text = "Leave Days"
    row = table.add_row().cells
    row[0].text = "HR"
    row[1].text = "20"
    document.add_paragraph("Requests should be submitted in advance.")
    document.save(path)


def test_pdf_parser_preserves_pages_and_text(tmp_path: Path) -> None:
    path = tmp_path / "policy.pdf"
    make_pdf(path)

    parsed = PdfParser().parse(path)

    assert parsed.page_count == 2
    assert [block.page_number for block in parsed.blocks] == [1, 2]
    assert "First page policy text." in parsed.blocks[0].content
    assert "Second page content." in parsed.full_text


def test_pdf_parser_rejects_corrupt_and_empty_text_files(tmp_path: Path) -> None:
    corrupt = tmp_path / "corrupt.pdf"
    corrupt.write_bytes(b"not a pdf")
    with pytest.raises(ParserError, match="could not be parsed"):
        PdfParser().parse(corrupt)

    empty_text = tmp_path / "scan.pdf"
    document = fitz.open()
    document.new_page()
    document.save(empty_text)
    document.close()
    with pytest.raises(ParserError, match="No extractable text"):
        PdfParser().parse(empty_text)


def test_docx_parser_preserves_order_and_tables_without_page_numbers(tmp_path: Path) -> None:
    path = tmp_path / "policy.docx"
    make_docx(path)

    parsed = DocxParser().parse(path)

    assert parsed.page_count is None
    assert all(block.page_number is None for block in parsed.blocks)
    assert [block.sequence_index for block in parsed.blocks] == [0, 1, 2, 3]
    assert "Leave Policy" in parsed.blocks[0].content
    assert "Department | Leave Days" in parsed.blocks[2].content
    assert "HR | 20" in parsed.blocks[2].content


def test_docx_parser_rejects_corrupt_file(tmp_path: Path) -> None:
    path = tmp_path / "corrupt.docx"
    path.write_bytes(b"not a docx")

    with pytest.raises(ParserError, match="could not be parsed"):
        DocxParser().parse(path)
